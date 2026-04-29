#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
测试文档知识库功能

测试通用文档扫描器 (txt/md/html/docx/网页)
"""

import sys
import os
import tempfile
import shutil
from pathlib import Path

project_root = Path(__file__).parent.parent
if str(project_root) not in sys.path:
    sys.path.insert(0, str(project_root))

from src.services.knowledge_base.scan_generic_documents import (
    GenericDocumentScanner,
    TextProcessor,
    MarkdownProcessor,
    HTMLProcessor,
    DocxProcessor,
    WebDocumentProcessor
)


def create_test_documents(test_dir: Path):
    """创建测试文档"""
    test_dir.mkdir(parents=True, exist_ok=True)
    
    (test_dir / "test.txt").write_text(
        "测试文本文件\n"
        "这是第一行内容\n"
        "这是第二行内容\n",
        encoding='utf-8'
    )
    
    (test_dir / "test.md").write_text(
        "# 测试 Markdown 文档\n\n"
        "## 简介\n\n"
        "这是一个测试文档。\n\n"
        "## 代码示例\n\n"
        "```delphi\n"
        "procedure Test;\n"
        "begin\n"
        "  ShowMessage('Hello');\n"
        "end;\n"
        "```\n",
        encoding='utf-8'
    )
    
    (test_dir / "test.html").write_text(
        "<!DOCTYPE html>\n"
        "<html>\n"
        "<head><title>测试 HTML 文档</title></head>\n"
        "<body>\n"
        "<h1>标题</h1>\n"
        "<p>这是段落内容。</p>\n"
        "<pre><code>procedure Test; begin end;</code></pre>\n"
        "</body>\n"
        "</html>\n",
        encoding='utf-8'
    )
    
    print(f"测试文档已创建: {test_dir}")


def test_text_processor():
    """测试文本处理器"""
    print("\n" + "=" * 60)
    print("测试 1: 文本处理器 (.txt)")
    print("=" * 60)
    
    processor = TextProcessor()
    
    with tempfile.TemporaryDirectory() as tmpdir:
        test_file = Path(tmpdir) / "test.txt"
        test_file.write_text("标题行\n内容行1\n内容行2\n", encoding='utf-8')
        
        assert processor.can_process(test_file), "应该能处理 .txt 文件"
        
        result = processor.process(test_file)
        
        assert result is not None, "处理结果不应为 None"
        assert result['content_type'] == 'text', "内容类型应为 text"
        assert result['title'] == '标题行', f"标题应为'标题行', 实际为'{result['title']}'"
        assert result['line_count'] >= 3, f"行数应至少为3, 实际为{result['line_count']}"
        
        print(f"  标题: {result['title']}")
        print(f"  内容类型: {result['content_type']}")
        print(f"  行数: {result['line_count']}")
        print(f"  大小: {result['size']} 字节")
        print("  ✓ 测试通过")
    
    return True


def test_markdown_processor():
    """测试 Markdown 处理器"""
    print("\n" + "=" * 60)
    print("测试 2: Markdown 处理器 (.md)")
    print("=" * 60)
    
    processor = MarkdownProcessor()
    
    with tempfile.TemporaryDirectory() as tmpdir:
        test_file = Path(tmpdir) / "test.md"
        test_file.write_text(
            "# 主标题\n\n"
            "## 章节1\n\n"
            "内容1\n\n"
            "```delphi\n"
            "procedure Test;\n"
            "```\n",
            encoding='utf-8'
        )
        
        assert processor.can_process(test_file), "应该能处理 .md 文件"
        
        result = processor.process(test_file)
        
        assert result is not None, "处理结果不应为 None"
        assert result['content_type'] == 'markdown', "内容类型应为 markdown"
        assert result['title'] == '主标题', f"标题应为'主标题', 实际为'{result['title']}'"
        assert len(result['sections']) >= 2, f"应至少有2个章节, 实际有{len(result['sections'])}"
        assert len(result['code_examples']) >= 1, f"应至少有1个代码块, 实际有{len(result['code_examples'])}"
        
        print(f"  标题: {result['title']}")
        print(f"  内容类型: {result['content_type']}")
        print(f"  章节数: {len(result['sections'])}")
        print(f"  代码块数: {len(result['code_examples'])}")
        print("  ✓ 测试通过")
    
    return True


def test_html_processor():
    """测试 HTML 处理器"""
    print("\n" + "=" * 60)
    print("测试 3: HTML 处理器 (.html)")
    print("=" * 60)
    
    processor = HTMLProcessor()
    
    with tempfile.TemporaryDirectory() as tmpdir:
        test_file = Path(tmpdir) / "test.html"
        test_file.write_text(
            "<!DOCTYPE html>\n"
            "<html><head><title>测试标题</title></head>\n"
            "<body>\n"
            "<h1>标题1</h1>\n"
            "<h2>标题2</h2>\n"
            "<pre><code>code here</code></pre>\n"
            "</body></html>\n",
            encoding='utf-8'
        )
        
        assert processor.can_process(test_file), "应该能处理 .html 文件"
        
        result = processor.process(test_file)
        
        assert result is not None, "处理结果不应为 None"
        assert result['content_type'] == 'html', "内容类型应为 html"
        assert '测试标题' in result['title'], f"标题应包含'测试标题', 实际为'{result['title']}'"
        
        print(f"  标题: {result['title']}")
        print(f"  内容类型: {result['content_type']}")
        print(f"  章节数: {len(result.get('sections', []))}")
        print(f"  代码块数: {len(result.get('code_examples', []))}")
        print("  ✓ 测试通过")
    
    return True


def test_docx_processor():
    """测试 Word 文档处理器"""
    print("\n" + "=" * 60)
    print("测试 4: Word 文档处理器 (.docx)")
    print("=" * 60)
    
    try:
        from docx import Document
    except ImportError:
        print("  ⚠ 跳过测试: python-docx 未安装")
        print("    安装命令: pip install python-docx")
        return True
    
    processor = DocxProcessor()
    
    with tempfile.TemporaryDirectory() as tmpdir:
        test_file = Path(tmpdir) / "test.docx"
        
        doc = Document()
        doc.add_heading('测试文档', level=1)
        doc.add_paragraph('这是第一段内容。')
        doc.add_paragraph('这是第二段内容。')
        doc.save(str(test_file))
        
        assert processor.can_process(test_file), "应该能处理 .docx 文件"
        
        result = processor.process(test_file)
        
        assert result is not None, "处理结果不应为 None"
        assert result['content_type'] == 'docx', "内容类型应为 docx"
        
        print(f"  标题: {result['title']}")
        print(f"  内容类型: {result['content_type']}")
        print(f"  行数: {result['line_count']}")
        print("  ✓ 测试通过")
    
    return True


def test_scanner_scan_directory():
    """测试扫描目录"""
    print("\n" + "=" * 60)
    print("测试 5: 扫描目录")
    print("=" * 60)
    
    with tempfile.TemporaryDirectory() as tmpdir:
        test_dir = Path(tmpdir) / "docs"
        create_test_documents(test_dir)
        
        db_path = Path(tmpdir) / "test.db"
        scanner = GenericDocumentScanner(str(db_path.parent), config={'database': {'file': 'test.db'}})
        
        result = scanner.scan_directory(str(test_dir))
        
        assert 'error' not in result, f"扫描不应出错: {result.get('error')}"
        assert result['total_files'] >= 3, f"应至少有3个文件, 实际有{result['total_files']}"
        assert result['processed'] >= 3, f"应至少处理成功3个文件, 实际处理{result['processed']}"
        
        print(f"  总文件数: {result['total_files']}")
        print(f"  处理成功: {result['processed']}")
        print(f"  处理失败: {result['failed']}")
        
        stats = scanner.get_statistics()
        print(f"  总文档数: {stats['total_documents']}")
        print(f"  按类型: {stats['by_type']}")
        print("  ✓ 测试通过")
    
    return True


def test_scanner_search():
    """测试搜索文档"""
    print("\n" + "=" * 60)
    print("测试 6: 搜索文档")
    print("=" * 60)
    
    with tempfile.TemporaryDirectory() as tmpdir:
        test_dir = Path(tmpdir) / "docs"
        create_test_documents(test_dir)
        
        db_path = Path(tmpdir) / "test.db"
        scanner = GenericDocumentScanner(str(db_path.parent), config={'database': {'file': 'test.db'}})
        
        scanner.scan_directory(str(test_dir))
        
        results = scanner.search("测试")
        
        assert len(results) > 0, "应能搜索到文档"
        
        print(f"  搜索 '测试' 找到 {len(results)} 个结果:")
        for i, doc in enumerate(results[:3], 1):
            print(f"    {i}. {doc['title']} ({doc['content_type']})")
        
        print("  ✓ 测试通过")
    
    return True


def test_scanner_content_type_filter():
    """测试按类型过滤搜索"""
    print("\n" + "=" * 60)
    print("测试 7: 按类型过滤搜索")
    print("=" * 60)
    
    with tempfile.TemporaryDirectory() as tmpdir:
        test_dir = Path(tmpdir) / "docs"
        create_test_documents(test_dir)
        
        db_path = Path(tmpdir) / "test.db"
        scanner = GenericDocumentScanner(str(db_path.parent), config={'database': {'file': 'test.db'}})
        
        scanner.scan_directory(str(test_dir))
        
        results = scanner.search("测试", content_type="markdown")
        
        print(f"  搜索类型=markdown 找到 {len(results)} 个结果:")
        for i, doc in enumerate(results[:3], 1):
            print(f"    {i}. {doc['title']} ({doc['content_type']})")
            assert doc['content_type'] == 'markdown', "结果类型应为 markdown"
        
        print("  ✓ 测试通过")
    
    return True


def test_web_processor():
    """测试网页处理器"""
    print("\n" + "=" * 60)
    print("测试 8: 网页处理器")
    print("=" * 60)
    
    try:
        import requests
    except ImportError:
        print("  ⚠ 跳过测试: requests 未安装")
        return True
    
    processor = WebDocumentProcessor()
    
    test_url = "https://example.com"
    print(f"  测试 URL: {test_url}")
    
    result = processor.process_url(test_url, timeout=10)
    
    if result is None:
        print("  ⚠ 网页抓取失败 (可能是网络问题)")
        return True
    
    assert 'title' in result, "结果应包含 title"
    assert 'content' in result, "结果应包含 content"
    
    print(f"  标题: {result['title']}")
    print(f"  内容类型: {result['content_type']}")
    print(f"  大小: {result['size']} 字节")
    print("  ✓ 测试通过")
    
    return True


def test_scanner_add_web_document():
    """测试添加网页文档"""
    print("\n" + "=" * 60)
    print("测试 9: 添加网页文档")
    print("=" * 60)
    
    try:
        import requests
    except ImportError:
        print("  ⚠ 跳过测试: requests 未安装")
        return True
    
    with tempfile.TemporaryDirectory() as tmpdir:
        db_path = Path(tmpdir) / "test.db"
        scanner = GenericDocumentScanner(str(db_path.parent), config={'database': {'file': 'test.db'}})
        
        test_url = "https://example.com"
        print(f"  添加网页: {test_url}")
        
        result = scanner.add_web_document(test_url)
        
        if result is None or (isinstance(result, dict) and 'error' in result):
            print("  ⚠ 添加网页失败 (可能是网络问题)")
            return True
        
        print(f"  标题: {result['title']}")
        print(f"  URL: {result.get('url', 'N/A')}")
        
        stats = scanner.get_statistics()
        print(f"  总文档数: {stats['total_documents']}")
        print("  ✓ 测试通过")
    
    return True


def main():
    """运行所有测试"""
    print("=" * 60)
    print("文档知识库功能测试")
    print("=" * 60)
    
    tests = [
        ("文本处理器", test_text_processor),
        ("Markdown处理器", test_markdown_processor),
        ("HTML处理器", test_html_processor),
        ("Word处理器", test_docx_processor),
        ("扫描目录", test_scanner_scan_directory),
        ("搜索文档", test_scanner_search),
        ("按类型过滤", test_scanner_content_type_filter),
        ("网页处理器", test_web_processor),
        ("添加网页", test_scanner_add_web_document),
    ]
    
    passed = 0
    failed = 0
    skipped = 0
    
    for name, test_func in tests:
        try:
            result = test_func()
            if result:
                passed += 1
            else:
                failed += 1
        except AssertionError as e:
            print(f"  ✗ 断言失败: {e}")
            failed += 1
        except Exception as e:
            print(f"  ✗ 测试异常: {e}")
            failed += 1
    
    print("\n" + "=" * 60)
    print("测试总结")
    print("=" * 60)
    print(f"通过: {passed}")
    print(f"失败: {failed}")
    print(f"跳过: {skipped}")
    print(f"总计: {len(tests)}")
    
    return failed == 0


if __name__ == "__main__":
    success = main()
    sys.exit(0 if success else 1)
