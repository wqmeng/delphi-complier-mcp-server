#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
通用文档扫描器
支持 doc/docx/txt/md/html/pdf/epub/hlp/网页等多种文档格式
"""

import re
import json
import struct
import hashlib
import sqlite3
import logging
import threading
import subprocess
import tempfile
import shutil
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Optional, Tuple, Callable
from concurrent.futures import ProcessPoolExecutor, as_completed
from multiprocessing import cpu_count

from .fts5_lazy_manager import FTS5LazyManager

logger = logging.getLogger(__name__)

# 默认排除的多语言子目录
DEFAULT_EXCLUDE_DIRS = {'ja', 'fr', 'de', 'es', 'it', 'ko', 'pt', 'ru', 'zh'}

try:
    from bs4 import BeautifulSoup
except ImportError:
    BeautifulSoup = None

try:
    import html2text
except ImportError:
    html2text = None


class DocumentProcessor:
    """文档处理器基类"""
    
    def __init__(self):
        self.supported_extensions = []
    
    def can_process(self, file_path: Path) -> bool:
        return file_path.suffix.lower() in self.supported_extensions
    
    def process(self, file_path: Path) -> Optional[Dict]:
        raise NotImplementedError


class TextProcessor(DocumentProcessor):
    """纯文本处理器 (.txt)"""
    
    def __init__(self):
        self.supported_extensions = ['.txt']
    
    def process(self, file_path: Path) -> Optional[Dict]:
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
            
            lines = content.split('\n')
            title = lines[0].strip()[:100] if lines else file_path.stem
            
            return {
                'title': title,
                'content': content,
                'content_type': 'text',
                'size': len(content),
                'line_count': len(lines),
                'hash': hashlib.md5(content.encode()).hexdigest(),
                'sections': [],
                'code_examples': []
            }
        except Exception:
            return None


class MarkdownProcessor(DocumentProcessor):
    """Markdown 处理器 (.md)"""
    
    def __init__(self):
        self.supported_extensions = ['.md', '.markdown']
    
    def process(self, file_path: Path) -> Optional[Dict]:
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
            
            title = self._extract_title(content) or file_path.stem
            sections = self._extract_sections(content)
            code_examples = self._extract_code_blocks(content)
            
            return {
                'title': title,
                'content': content,
                'content_type': 'markdown',
                'size': len(content),
                'line_count': content.count('\n') + 1,
                'hash': hashlib.md5(content.encode()).hexdigest(),
                'sections': sections,
                'code_examples': code_examples
            }
        except Exception:
            return None
    
    def _extract_title(self, content: str) -> Optional[str]:
        for line in content.split('\n'):
            if line.startswith('#'):
                return line.lstrip('#').strip()[:100]
        return None
    
    def _extract_sections(self, content: str) -> List[Dict]:
        sections = []
        for line in content.split('\n'):
            if line.startswith('#'):
                level = len(line) - len(line.lstrip('#'))
                title = line.lstrip('#').strip()
                sections.append({'level': level, 'title': title})
        return sections
    
    def _extract_code_blocks(self, content: str) -> List[str]:
        pattern = re.compile(r'```[\w]*\n(.*?)```', re.DOTALL)
        return [m.group(1).strip() for m in pattern.finditer(content)]


class HTMLProcessor(DocumentProcessor):
    """HTML 处理器 (.htm, .html)"""
    
    def __init__(self):
        self.supported_extensions = ['.htm', '.html']
        if html2text:
            self.converter = html2text.HTML2Text()
            self.converter.ignore_links = False
            self.converter.ignore_images = True
            self.converter.body_width = 0
        else:
            self.converter = None
    
    def process(self, file_path: Path) -> Optional[Dict]:
        if not BeautifulSoup:
            return None
        
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                html_content = f.read()
            
            soup = BeautifulSoup(html_content, 'lxml' if _has_lxml() else 'html.parser')
            
            for tag in soup(['script', 'style', 'nav', 'footer']):
                tag.decompose()
            
            title = soup.title.get_text(strip=True) if soup.title else file_path.stem
            text_content = soup.get_text(separator='\n', strip=True)
            
            if self.converter:
                markdown = self.converter.handle(html_content)
            else:
                markdown = text_content
            
            sections = self._extract_sections(soup)
            code_examples = self._extract_code_blocks(soup)
            
            return {
                'title': title,
                'content': markdown,
                'content_type': 'html',
                'size': len(text_content),
                'line_count': text_content.count('\n') + 1,
                'hash': hashlib.md5(text_content.encode()).hexdigest(),
                'sections': sections,
                'code_examples': code_examples
            }
        except Exception:
            return None
    
    def _extract_sections(self, soup) -> List[Dict]:
        sections = []
        for tag in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6']):
            level = int(tag.name[1])
            title = tag.get_text(strip=True)
            sections.append({'level': level, 'title': title})
        return sections
    
    def _extract_code_blocks(self, soup) -> List[str]:
        return [tag.get_text(strip=True) for tag in soup.find_all(['pre', 'code'])]


class DocxProcessor(DocumentProcessor):
    """Word 文档处理器 (.docx)"""
    
    def __init__(self):
        self.supported_extensions = ['.docx']
        self.has_docx = self._check_docx()
    
    def _check_docx(self) -> bool:
        try:
            from docx import Document
            return True
        except ImportError:
            return False
    
    def process(self, file_path: Path) -> Optional[Dict]:
        if not self.has_docx:
            return {
                'title': file_path.stem,
                'content': f"Word .docx 处理需要安装依赖:\n"
                          f"  pip install python-docx",
                'content_type': 'docx',
                'size': 0,
                'line_count': 0,
                'hash': '',
                'sections': [],
                'code_examples': [],
                'requires_dependencies': True,
                'install_hint': 'pip install python-docx'
            }
        
        try:
            from docx import Document
            doc = Document(str(file_path))
            
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            content = '\n'.join(paragraphs)
            
            title = paragraphs[0][:100] if paragraphs else file_path.stem
            sections = self._extract_sections(doc)
            code_examples = self._extract_code_examples(paragraphs)
            
            return {
                'title': title,
                'content': content,
                'content_type': 'docx',
                'size': len(content),
                'line_count': len(paragraphs),
                'hash': hashlib.md5(content.encode()).hexdigest(),
                'sections': sections,
                'code_examples': code_examples
            }
        except Exception:
            return None
    
    def _extract_sections(self, doc) -> List[Dict]:
        sections = []
        for p in doc.paragraphs:
            if p.style.name.startswith('Heading'):
                try:
                    level = int(p.style.name.split()[-1])
                except (ValueError, IndexError):
                    level = 1
                sections.append({'level': level, 'title': p.text})
        return sections
    
    def _extract_code_examples(self, paragraphs: List[str]) -> List[str]:
        code_blocks = []
        in_code = False
        current_block = []
        
        for p in paragraphs:
            if '```' in p or p.strip().startswith('    '):
                if not in_code:
                    in_code = True
                    current_block = []
                current_block.append(p)
            elif in_code:
                code_blocks.append('\n'.join(current_block))
                in_code = False
                current_block = []
        
        return code_blocks


class DocProcessor(DocumentProcessor):
    """旧版 Word 文档处理器 (.doc) - 需要 antiword 或 catdoc"""
    
    def __init__(self):
        self.supported_extensions = ['.doc']
    
    def process(self, file_path: Path) -> Optional[Dict]:
        try:
            
            result = subprocess.run(
                ['antiword', str(file_path)],
                capture_output=True,
                text=True,
                encoding='utf-8',
                errors='ignore',
                timeout=30,
                creationflags=getattr(subprocess, 'CREATE_NO_WINDOW', 0)
            )
            
            if result.returncode == 0:
                content = result.stdout
            else:
                result = subprocess.run(
                    ['catdoc', '-w', str(file_path)],
                    capture_output=True,
                    text=True,
                    encoding='utf-8',
                    errors='ignore',
                    timeout=30,
                    creationflags=getattr(subprocess, 'CREATE_NO_WINDOW', 0)
                )
                if result.returncode != 0:
                    return None
                content = result.stdout
            
            lines = content.split('\n')
            title = lines[0].strip()[:100] if lines else file_path.stem
            
            return {
                'title': title,
                'content': content,
                'content_type': 'doc',
                'size': len(content),
                'line_count': len(lines),
                'hash': hashlib.md5(content.encode()).hexdigest(),
                'sections': [],
                'code_examples': []
            }
        except Exception:
            return None


class PDFProcessor(DocumentProcessor):
    """PDF 文档处理器 (.pdf) - 需要 pdfplumber 或 PyMuPDF"""
    
    def __init__(self):
        self.supported_extensions = ['.pdf']
        self.has_pdfplumber = self._check_pdfplumber()
        self.has_pymupdf = self._check_pymupdf()
    
    def _check_pdfplumber(self) -> bool:
        try:
            import pdfplumber
            return True
        except ImportError:
            return False
    
    def _check_pymupdf(self) -> bool:
        try:
            import fitz
            return True
        except ImportError:
            return False
    
    def process(self, file_path: Path) -> Optional[Dict]:
        # 优先使用 PyMuPDF（性能更好）
        if self.has_pymupdf:
            return self._process_with_pymupdf(file_path)
        
        # 备选：使用 pdfplumber
        if self.has_pdfplumber:
            return self._process_with_pdfplumber(file_path)
        
        # 都没有安装，返回提示信息
        return {
            'title': file_path.stem,
            'content': f"PDF 处理需要安装依赖（推荐优先使用 PyMuPDF）:\n"
                      f"  pip install PyMuPDF\n"
                      f"  或（备选）\n"
                      f"  pip install pdfplumber",
            'content_type': 'pdf',
            'size': 0,
            'line_count': 0,
            'hash': '',
            'sections': [],
            'code_examples': [],
            'requires_dependencies': True,
            'install_hint': 'pip install PyMuPDF'
        }
    
    def _process_with_pymupdf(self, file_path: Path) -> Optional[Dict]:
        """使用 PyMuPDF 处理（高性能）"""
        try:
            import fitz
            
            doc = fitz.open(str(file_path))
            
            text_content = []
            sections = []
            
            for page_num, page in enumerate(doc):
                text = page.get_text()
                text_content.append(text)
                
                # 提取标题（基于字体大小）
                blocks = page.get_text("dict")["blocks"]
                for block in blocks:
                    if "lines" in block:
                        for line in block["lines"]:
                            for span in line["spans"]:
                                if span["size"] > 14:
                                    title = span["text"].strip()
                                    if title and len(title) < 100:
                                        sections.append({'level': 1, 'title': title})
            
            full_text = '\n'.join(text_content)
            metadata = doc.metadata
            title = metadata.get('title') or file_path.stem
            
            doc.close()
            
            return {
                'title': title,
                'content': full_text,
                'content_type': 'pdf',
                'size': len(full_text),
                'line_count': full_text.count('\n') + 1,
                'hash': hashlib.md5(full_text.encode()).hexdigest(),
                'sections': sections[:50],
                'code_examples': [],
                'metadata': {
                    'author': metadata.get('author'),
                    'subject': metadata.get('subject'),
                    'page_count': len(text_content)
                }
            }
        except Exception:
            return None
    
    def _process_with_pdfplumber(self, file_path: Path) -> Optional[Dict]:
        """使用 pdfplumber 处理"""
        try:
            import pdfplumber
            
            text_content = []
            
            with pdfplumber.open(str(file_path)) as pdf:
                metadata = pdf.metadata or {}
                title = metadata.get('Title') or file_path.stem
                
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        text_content.append(text)
                
                full_text = '\n'.join(text_content)
                
                return {
                    'title': title,
                    'content': full_text,
                    'content_type': 'pdf',
                    'size': len(full_text),
                    'line_count': full_text.count('\n') + 1,
                    'hash': hashlib.md5(full_text.encode()).hexdigest(),
                    'sections': [],
                    'code_examples': [],
                    'metadata': {
                        'author': metadata.get('Author'),
                        'page_count': len(pdf.pages)
                    }
                }
        except Exception:
            return None


class EPUBProcessor(DocumentProcessor):
    """EPUB 电子书处理器 - 使用 zipfile + BeautifulSoup 直接解析"""
    
    def __init__(self):
        self.supported_extensions = ['.epub']
    
    def process(self, file_path: Path) -> Optional[Dict]:
        try:
            import zipfile
            from xml.etree import ElementTree as ET
            
            with zipfile.ZipFile(str(file_path), 'r') as zf:
                container_xml = zf.read('META-INF/container.xml')
                container_root = ET.fromstring(container_xml)
                
                rootfile_elem = container_root.find('.//{urn:oasis:names:tc:opendocument:xmlns:container}rootfile')
                if rootfile_elem is None:
                    return None
                
                opf_path = rootfile_elem.get('full-path')
                if not opf_path:
                    return None
                
                opf_content = zf.read(opf_path)
                opf_root = ET.fromstring(opf_content)
                
                ns = {'opf': 'http://www.idpf.org/2007/opf'}
                
                manifest = {}
                for item in opf_root.findall('.//opf:item', ns):
                    item_id = item.get('id')
                    href = item.get('href')
                    if item_id and href:
                        manifest[item_id] = href
                
                spine_ids = []
                for itemref in opf_root.findall('.//opf:itemref', ns):
                    idref = itemref.get('idref')
                    if idref in manifest:
                        spine_ids.append(idref)
                
                opf_dir = opf_path.rsplit('/', 1)[0] if '/' in opf_path else ''
                
                text_content = []
                sections = []
                
                for item_id in spine_ids[:100]:
                    href = manifest[item_id]
                    if opf_dir:
                        file_path_in_zip = f"{opf_dir}/{href}"
                    else:
                        file_path_in_zip = href
                    
                    try:
                        html_content = zf.read(file_path_in_zip)
                        
                        if BeautifulSoup:
                            soup = BeautifulSoup(html_content, 'lxml')
                            
                            for heading in soup.find_all(['h1', 'h2', 'h3']):
                                title = heading.get_text().strip()
                                if title and len(title) < 100:
                                    level = int(heading.name[1])
                                    sections.append({'level': level, 'title': title})
                            
                            for script in soup.find_all('script'):
                                script.decompose()
                            for style in soup.find_all('style'):
                                style.decompose()
                            
                            text = soup.get_text(separator='\n', strip=True)
                            if text:
                                text_content.append(text)
                        else:
                            text = html_content.decode('utf-8', errors='ignore')
                            text_content.append(text)
                    except KeyError:
                        continue
                
                full_text = '\n\n'.join(text_content)
                
                title = file_path.stem
                for elem in opf_root.findall('.//opf:title', ns):
                    if elem.text:
                        title = elem.text
                        break
                
                metadata = {}
                for elem in opf_root.findall('.//opf:creator', ns):
                    if elem.text:
                        metadata['author'] = elem.text
                        break
                
                return {
                    'title': title,
                    'content': full_text,
                    'content_type': 'epub',
                    'size': len(full_text),
                    'line_count': full_text.count('\n') + 1,
                    'hash': hashlib.md5(full_text.encode()).hexdigest(),
                    'sections': sections[:50],
                    'code_examples': [],
                    'metadata': metadata
                }
        except Exception:
            return None


class ChmProcessor(DocumentProcessor):
    """Microsoft Compiled HTML Help 处理器 (.chm).

    使用 7-Zip 解压 CHM 文件后，将内部的 HTML 文件作为文档导入。
    需要系统安装 7-Zip (https://www.7-zip.org/)。
    """

    # 7z 解压时跳过的辅助文件类型
    EXCLUDE_PATTERNS = [
        '-x!*.png', '-x!*.gif', '-x!*.jpg', '-x!*.jpeg',
        '-x!*.css', '-x!*.js', '-x!*.ico', '-x!*.svg',
        '-x!*.woff', '-x!*.woff2', '-x!*.ttf', '-x!*.eot',
        '-x!*.bmp', '-x!*.webp', '-x!*.avif',
    ]

    def __init__(self):
        self.supported_extensions = ['.chm']
        self._sevenzip_path = None

    def _find_7zip(self) -> Optional[str]:
        if self._sevenzip_path:
            return self._sevenzip_path
        self._sevenzip_path = self._find_7zip_path()
        return self._sevenzip_path

    def process(self, file_path: Path) -> Optional[List[Dict]]:
        sevenzip = self._find_7zip()
        if not sevenzip:
            return [{
                'title': file_path.stem,
                'content': (
                    '需要安装 7-Zip 才能处理 CHM 文件\n\n'
                    '下载地址:\n'
                    '  官网: https://www.7-zip.org/download.html\n'
                    '  SourceForge: https://sourceforge.net/projects/sevenzip/files/7-Zip/\n\n'
                    'Windows 64位用户下载 7z2301-x64.exe 安装即可'
                ),
                'content_type': 'chm',
                'size': 0, 'line_count': 0, 'hash': '', 'sections': [], 'code_examples': [],
                'requires_dependencies': True,
                'install_hint': '需要安装 7-Zip: https://www.7-zip.org/download.html'
            }]

        tmpdir = None
        try:
            # 用 7z 解压到临时目录（跳过图片/CSS/JS等辅助文件，只保留 HTML）
            tmpdir = tempfile.mkdtemp(prefix='chm_')
            result = subprocess.run(
                [sevenzip, 'x', '-y', f'-o{tmpdir}', str(file_path), *self.EXCLUDE_PATTERNS],
                stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL,
                timeout=600,
                creationflags=getattr(subprocess, 'CREATE_NO_WINDOW', 0)
            )
            if result.returncode != 0:
                return None

            # 搜集所有 htm/html 文件
            extracted_dir = Path(tmpdir)
            html_files = list(extracted_dir.rglob('*.htm')) + list(extracted_dir.rglob('*.html'))
            if not html_files:
                return None

            # 用 HTMLProcessor 逐一处理
            html_proc = HTMLProcessor()
            docs = []
            seen_hashes = set()
            for hf in html_files:
                try:
                    doc = html_proc.process(hf)
                    if doc:
                        rel_path = str(hf.relative_to(extracted_dir)).replace('\\', '/')
                        doc['path'] = rel_path
                        doc['full_path'] = str(hf)
                        doc['extension'] = '.html'
                        doc['file_size'] = hf.stat().st_size
                        doc['last_modified'] = datetime.fromtimestamp(
                            hf.stat().st_mtime
                        ).isoformat()
                        # 去重
                        doc_hash = doc.get('hash', '')
                        if doc_hash and doc_hash in seen_hashes:
                            continue
                        if doc_hash:
                            seen_hashes.add(doc_hash)
                        docs.append(doc)
                except Exception:
                    continue

            return docs if docs else None

        except Exception:
            return None
        finally:
            if tmpdir:
                shutil.rmtree(tmpdir, ignore_errors=True)


class HlpProcessor(DocumentProcessor):
    """Windows Help 文件处理器 (.hlp) — 纯 Python, 输出 Markdown.
    解析 |TTLBTREE 获得标题, |TOPIC 解压后通过 Hall/Phrase 解码获得正文.
    支持 HC30/HC31/HCW 4.00 格式, LZ77 压缩.
    """

    def __init__(self):
        self.supported_extensions = ['.hlp']

    # ═══════════════════════════════════════════
    #  公开入口 — 按 RT2 topic header 拆分为多个文档
    # ═══════════════════════════════════════════

    def process(self, file_path: Path) -> Optional[List[Dict]]:
        try:
            with open(file_path, 'rb') as f:
                header = f.read(16)
                if len(header) < 16:
                    return None
                magic, dir_start, _fb, _fs = struct.unpack('<LLLl', header)
                if magic != 0x00035F3F:
                    return None

                internal_files = self._parse_directory(f, dir_start)
                if not internal_files or '|SYSTEM' not in internal_files:
                    return None

                system_info = self._read_system(f, internal_files)
                if not system_info:
                    return None

                phrases = self._parse_phrases(f, internal_files, system_info)
                titles = self._read_ttlbtree(f, internal_files)
                topic_list = self._extract_topic_text(f, internal_files, system_info, phrases, titles)

        except Exception:
            return None

        if not topic_list:
            return None

        docs = []
        for title, body_text in topic_list:
            if not title and not body_text:
                continue
            content_parts = [f'## {title}'] if title else []
            if body_text:
                content_parts.append(body_text)
            content = '\n\n'.join(content_parts)
            content = re.sub(r'\n{3,}', '\n\n', content).strip()
            if not content:
                continue

            sections = [{'level': 2, 'title': title}] if title else []
            code_examples = re.findall(r'```[\w]*\n(.*?)```', content, re.DOTALL)

            docs.append({
                'title': (title or file_path.stem)[:200],
                'content': content,
                'content_type': 'hlp',
                'size': len(content),
                'line_count': content.count('\n') + 1,
                'hash': hashlib.md5(content.encode()).hexdigest(),
                'sections': sections,
                'code_examples': code_examples
            })

        return docs if docs else None

    # ═══════════════════════════════════════════
    #  目录 B+ 树
    # ═══════════════════════════════════════════

    def _parse_directory(self, f, dir_start):
        f.seek(dir_start + 9)
        f.read(2); f.read(2)
        page_size = struct.unpack('<H', f.read(2))[0]
        f.read(16); f.read(2); f.read(2)
        root_page = struct.unpack('<h', f.read(2))[0]
        f.read(2); f.read(2)
        nlevels = struct.unpack('<h', f.read(2))[0]
        f.read(4)

        def read_leaf_fn(fh, n):
            r = []
            for _ in range(n):
                nb = bytearray()
                while True:
                    b = fh.read(1)
                    if b in (b'\x00', b''): break
                    nb.extend(b)
                r.append((nb.decode('latin-1', errors='replace'),
                          struct.unpack('<L', fh.read(4))[0]))
            return r

        page_start = dir_start + 9 + 38
        entries = self._traverse_btree_leaves(f, page_start, root_page, nlevels, page_size, read_leaf_fn)
        return dict(entries)

    def _traverse_btree_leaves(self, f, page_start, root_page, nlevels, page_size, read_leaf_fn):
        page_num = root_page
        for _ in range(nlevels - 1):
            f.seek(page_start + page_num * page_size)
            f.read(2); nentries = struct.unpack('<h', f.read(2))[0]
            prev_page = struct.unpack('<h', f.read(2))[0]
            names = []
            for __ in range(nentries):
                nb = bytearray()
                while True:
                    b = f.read(1)
                    if b in (b'\x00', b''): break
                    nb.extend(b)
                names.append(struct.unpack('<h', f.read(2))[0])
            page_num = prev_page if prev_page != -1 else (names[0] if names else -1)

        results = []
        while page_num not in (-1, 0xFFFF):
            f.seek(page_start + page_num * page_size)
            f.read(2); nentries = struct.unpack('<h', f.read(2))[0]
            f.read(2); next_page = struct.unpack('<h', f.read(2))[0]
            results.extend(read_leaf_fn(f, nentries))
            page_num = next_page
        return results

    # ═══════════════════════════════════════════
    #  |SYSTEM
    # ═══════════════════════════════════════════

    def _read_system(self, f, internal_files):
        sys_offset = internal_files.get('|SYSTEM')
        if sys_offset is None:
            return None
        f.seek(sys_offset); f.read(4); f.read(4); f.read(1)
        data_start = sys_offset + 9
        f.seek(data_start)
        magic = struct.unpack('<H', f.read(2))[0]
        if magic != 0x036C: return None
        minor = struct.unpack('<h', f.read(2))[0]
        f.read(2); f.read(4)
        flags = struct.unpack('<H', f.read(2))[0]

        tb_size = 4096 if (minor > 16 and (flags & 4)) else 2048
        compressed = minor > 16 and (flags & 0xC)
        has_hall = ('|PhrIndex' in internal_files) and ('|PhrImage' in internal_files)
        has_old_phrases = '|Phrases' in internal_files
        return {'topic_block_size': tb_size, 'compressed': compressed,
                'minor': minor, 'has_hall': has_hall, 'has_old_phrases': has_old_phrases}

    # ═══════════════════════════════════════════
    #  短语表 (Hall / old-style)
    # ═══════════════════════════════════════════

    def _parse_phrases(self, f, internal_files, system_info):
        if system_info.get('has_hall'):
            return self._parse_phrases_hall(f, internal_files)
        if system_info.get('has_old_phrases'):
            return self._parse_phrases_old(f, internal_files)
        return None

    # ── Hall (|PhrIndex + |PhrImage) ──

    def _parse_phrases_hall(self, f, internal_files):
        idx_off = internal_files.get('|PhrIndex')
        img_off = internal_files.get('|PhrImage')
        if not idx_off or not img_off:
            return None

        f.seek(idx_off); f.read(4); idx_used = struct.unpack('<L', f.read(4))[0]; f.read(1)
        idx_start = idx_off + 9
        f.seek(idx_start)
        idx_data = f.read(idx_used)

        magic = struct.unpack_from('<L', idx_data, 0)[0]
        if magic != 1: return None
        nentries = struct.unpack_from('<L', idx_data, 4)[0]
        img_size = struct.unpack_from('<L', idx_data, 12)[0]
        img_comp_size = struct.unpack_from('<L', idx_data, 16)[0]
        flags = struct.unpack_from('<H', idx_data, 24)[0]
        bit_count = flags & 0xF

        offsets = self._build_phrase_offsets(idx_data, 28, nentries, bit_count)
        if not offsets or len(offsets) < nentries + 1:
            return None

        f.seek(img_off); f.read(4); img_used = struct.unpack('<L', f.read(4))[0]; f.read(1)
        img_start = img_off + 9
        f.seek(img_start)
        img_raw = f.read(img_used)

        if img_comp_size != img_size:
            img_data = self._lz77_decompress(img_raw, img_size)
        else:
            img_data = img_raw

        phrases = []
        for i in range(nentries):
            s, e = offsets[i], offsets[i + 1]
            if s < len(img_data) and e <= len(img_data):
                phrases.append(img_data[s:e].decode('latin-1', errors='replace'))
            else:
                phrases.append('')
        return phrases

    def _build_phrase_offsets(self, data, start_offset, nentries, bit_count):
        byte_data = bytearray(data)
        ptr_idx = start_offset
        mask = 0
        value = 0

        def get_bit():
            nonlocal ptr_idx, mask, value
            mask = (mask << 1) & 0xFFFFFFFF  # Python ints don't overflow
            if mask == 0:
                if ptr_idx + 4 <= len(byte_data):
                    value = struct.unpack_from('<L', byte_data, ptr_idx)[0]
                else:
                    value = 0
                ptr_idx += 4
                mask = 1
            return (value & mask) != 0

        offsets = [0]
        for _ in range(nentries):
            n = 1
            while get_bit():
                n += 1 << bit_count
            if get_bit(): n += 1
            if bit_count > 1 and get_bit(): n += 2
            if bit_count > 2 and get_bit(): n += 4
            if bit_count > 3 and get_bit(): n += 8
            if bit_count > 4 and get_bit(): n += 16
            offsets.append(offsets[-1] + n)
        return offsets

    # ── old-style (|Phrases) ──

    def _parse_phrases_old(self, f, internal_files):
        phr_offset = internal_files.get('|Phrases')
        if phr_offset is None: return None
        f.seek(phr_offset); f.read(4); phr_used = struct.unpack('<L', f.read(4))[0]; f.read(1)
        phr_start = phr_offset + 9
        f.seek(phr_start)
        data = f.read(phr_used)
        if len(data) < 4: return None

        num_phrases = struct.unpack_from('<H', data, 0)[0]
        one_hundred = struct.unpack_from('<H', data, 2)[0]
        offsets = []
        pos = 4
        for _ in range(num_phrases + 1):
            offsets.append(struct.unpack_from('<H', data, pos)[0])
            pos += 2

        hdr_end = 4 + (num_phrases + 1) * 2
        if one_hundred == 0x0100 and len(data) > hdr_end + 4:
            decomp_size = struct.unpack_from('<L', data, hdr_end)[0]
            lz77_data = data[hdr_end + 4:]
            phrase_bytes = self._lz77_decompress(lz77_data, max(decomp_size, 65536))
            if not phrase_bytes:
                return None
        else:
            phrase_bytes = data[hdr_end:]

        phrases = []
        for i in range(num_phrases):
            s, e = offsets[i], offsets[i + 1]
            if s < len(phrase_bytes) and e <= len(phrase_bytes):
                phrases.append(phrase_bytes[s:e].decode('latin-1', errors='replace'))
            else:
                phrases.append('')
        return phrases

    # ═══════════════════════════════════════════
    #  |TOPIC 文本提取 (含 Hall / Phrase 解压)
    # ═══════════════════════════════════════════

    def _extract_topic_text(self, f, internal_files, system_info, phrases, titles):
        """Walk TOPICLINKs, group text by RT2 topic headers. Returns [(title, text)]."""
        topic_offset = internal_files.get('|TOPIC')
        if topic_offset is None:
            return []

        f.seek(topic_offset); f.read(4)
        topic_used = struct.unpack('<L', f.read(4))[0]; f.read(1)
        data_start = topic_offset + 9

        tb_size = system_info['topic_block_size']
        compressed = system_info['compressed']
        has_phrases = phrases is not None
        use_hall = system_info.get('has_hall', False)

        all_data = bytearray()
        offset = data_start
        while offset < data_start + topic_used:
            f.seek(offset); f.read(12)
            raw = f.read(tb_size - 12)
            if compressed:
                all_data.extend(self._lz77_decompress(raw, 16384))
            else:
                all_data.extend(raw)
            offset += tb_size

        topics = []  # [(title, [text_runs])]
        current_title = 'Unknown'
        current_texts = []

        title_index = 0
        pos = 0
        while pos + 21 <= len(all_data):
            bsz = struct.unpack_from('<L', all_data, pos)[0]
            if 10 <= bsz <= 4000:
                rt = all_data[pos + 20]
                dl2 = struct.unpack_from('<L', all_data, pos + 4)[0]
                dl1 = struct.unpack_from('<L', all_data, pos + 16)[0]

                if 21 < dl1 < bsz:
                    compressed_sz = bsz - dl1
                    ld2_raw = all_data[pos + dl1: pos + bsz]

                    ld2 = b''
                    if compressed_sz > 0 and dl2 > 0:
                        if has_phrases:
                            if dl2 <= compressed_sz:
                                ld2 = ld2_raw[:dl2]
                            elif use_hall:
                                ld2 = self._hall_decompress_bytes(ld2_raw, phrases, dl2)
                            else:
                                ld2 = self._phrase_decompress_bytes(ld2_raw, phrases, dl2)
                        else:
                            ld2 = ld2_raw

                    if rt == 2:
                        # Save previous topic
                        if title_index > 0:
                            topics.append((current_title, current_texts))
                            current_texts = []

                        # Extract title from LinkData2 first NUL-term string
                        title = self._extract_topic_header_text(ld2) if ld2 else ''
                        # Fallback to TTLBTREE title
                        if not title and title_index < len(titles):
                            title = titles[title_index][1]
                        current_title = title or f'Topic {title_index}'
                        title_index += 1

                    elif rt in (0x20, 0x23, 0x27) and ld2:
                        ld1 = all_data[pos + 21: pos + dl1]
                        t = self._extract_display_text(ld1, ld2)
                        if t:
                            words = t.split()
                            ratio = sum(1 for w in words if len(w) > 2) / max(len(words), 1)
                            if ratio > 0.3 or len(t) > 40:
                                current_texts.append(t)

                    pos += bsz
                    continue
            pos += 1

        # Last topic
        if title_index > 1 or (not topics and current_texts):
            topics.append((current_title, current_texts))

        # Merge text runs per topic
        result = []
        for i, (t_title, texts) in enumerate(topics):
            # Use TTLBTREE title when available (cleaner)
            if i < len(titles):
                t_title = titles[i][1]
            merged = '\n\n'.join(texts)
            merged = re.sub(r'\n{3,}', '\n\n', merged)
            result.append((t_title, merged))

        return result

    # ── RT=2 topic header: LinkData2 = NUL-separated title + macros ──

    @staticmethod
    def _extract_topic_header_text(ld2):
        """LinkData2 for RecordType 2: first NUL-terminated string is topic title."""
        parts = ld2.split(b'\x00')
        texts = []
        for part in parts:
            s = part.decode('latin-1', errors='replace')
            s = HlpProcessor._clean_text(s)
            if s:
                texts.append(s)
        return ' '.join(texts)

    # ── RT=0x20 display paragraph: LinkData1 has format header, LinkData2 = NUL-separated text ──

    @staticmethod
    def _extract_display_text(ld1, ld2):
        """Extract text from a display paragraph (RecordType 0x20/0x23/0x27).
        LinkData1 starts with format info; LinkData2 is NUL-terminated strings
        alternating with format control bytes from LinkData1."""
        off = 0

        # Skip expanded size (compressed long)
        off = HlpProcessor._skip_compressed_long(ld1, off)
        # Skip topic offset increment (word)
        off += 2

        # Parse ParagraphInfo flags (word)
        if off + 2 > len(ld1):
            return ''
        x2 = struct.unpack_from('<H', ld1, off)[0]
        off += 2

        # Skip conditional fields based on flags
        if x2 & 0x0001: off = HlpProcessor._skip_compressed_long(ld1, off)
        if x2 & 0x0002: off = HlpProcessor._skip_int(ld1, off)
        if x2 & 0x0004: off = HlpProcessor._skip_int(ld1, off)
        if x2 & 0x0008: off = HlpProcessor._skip_int(ld1, off)
        if x2 & 0x0010: off = HlpProcessor._skip_int(ld1, off)
        if x2 & 0x0020: off = HlpProcessor._skip_int(ld1, off)
        if x2 & 0x0040: off = HlpProcessor._skip_int(ld1, off)
        if x2 & 0x0100: off += 4  # border info (byte + word)
        if x2 & 0x0200:
            if off >= len(ld1): return ''
            ntabs = HlpProcessor._scan_int(ld1, off)
            off = HlpProcessor._skip_int(ld1, off)
            for _ in range(ntabs):
                if off + 2 > len(ld1): break
                tab = struct.unpack_from('<H', ld1, off)[0]
                off += 2
                if tab & 0x4000:
                    off += 2

        # Now iterate: read NUL-term string from ld2, then format byte from ld1.
        # Build paragraphs: LF(0x81)/CR(0x82)→soft break, 0xFF→paragraph end.
        paragraphs = []
        current = []
        ld2_pos = 0
        while ld2_pos < len(ld2):
            end = ld2.find(b'\x00', ld2_pos)
            if end < 0:
                end = len(ld2)
            s = ld2[ld2_pos:end].decode('latin-1', errors='replace')
            s = HlpProcessor._clean_text(s)
            if s:
                current.append(s)
            ld2_pos = end + 1

            if off >= len(ld1):
                break
            ctrl = ld1[off]
            if ctrl == 0xFF:
                if current:
                    paragraphs.append(' '.join(current))
                    current = []
                off += 1
                break  # end of paragraph
            elif ctrl in (0x81, 0x82):           # LF / CR → soft break
                if current:
                    paragraphs.append(' '.join(current))
                    current = []
                off += 1
            elif ctrl == 0x80: off += 3          # font change
            elif ctrl == 0x20: off += 5          # variable field
            elif ctrl == 0x21: off += 3          # dtype
            elif ctrl == 0x83: off += 4
            elif ctrl == 0x84: off += 5
            elif ctrl == 0x85: off += 6
            elif ctrl == 0x86: off += 7
            elif ctrl == 0x87: off += 8
            elif ctrl == 0x88: off += 9
            elif ctrl == 0x89: off += 10
            elif ctrl == 0x8A: off += 11
            elif ctrl == 0x8B: off += 12
            elif ctrl == 0x8C: off += 13
            elif ctrl == 0x23: off += 1          # topic separator
            else: off += 1

        if current:
            paragraphs.append(' '.join(current))
        return '\n\n'.join(paragraphs)

    # ── Compressed number helpers ──

    @staticmethod
    def _skip_compressed_long(data, off):
        while off < len(data) and (data[off] & 0x80):
            off += 1
        return off + 1

    @staticmethod
    def _skip_int(data, off):
        b = data[off]
        if b & 0x80:
            return off + 2
        return off + 1

    @staticmethod
    def _scan_int(data, off):
        b = data[off]
        if b & 0x80:
            return (b & 0x7F) | (data[off + 1] << 7)
        return b

    # ── Hall LinkData2 解压 ──

    @staticmethod
    def _hall_decompress(data, phrases, expected_len):
        return HlpProcessor._hall_decompress_bytes(data, phrases, expected_len).decode('latin-1', errors='replace')

    @staticmethod
    def _hall_decompress_bytes(data, phrases, expected_len):
        """Hall decompression, returns raw bytes (NUL-separated text fragments)."""
        result = bytearray()
        i = 0
        data_bytes = bytearray(data)
        while i < len(data_bytes):
            ch = data_bytes[i]
            if ch & 15 == 15:  i += 1
            elif ch & 15 == 7: result.append(0x20); i += 1
            elif ch & 7 == 3:
                copy_len = (ch >> 3) + 1; i += 1
                end = min(i + copy_len, len(data_bytes))
                result.extend(data_bytes[i:end])
                i = end
            elif ch & 3 == 1:
                if i + 1 >= len(data_bytes): break
                phrase_num = ch * 64 + 64 + data_bytes[i + 1]; i += 2
                if 0 <= phrase_num < len(phrases):
                    result.extend(phrases[phrase_num].encode('latin-1', errors='replace'))
            elif ch % 2 == 0:
                phrase_num = ch // 2; i += 1
                if 0 <= phrase_num < len(phrases):
                    result.extend(phrases[phrase_num].encode('latin-1', errors='replace'))
            else: i += 1
        return bytes(result[:expected_len])

    # ── old-style Phrase LinkData2 解压 ──

    @staticmethod
    def _phrase_decompress(data, phrases, expected_len):
        return HlpProcessor._phrase_decompress_bytes(data, phrases, expected_len).decode('latin-1', errors='replace')

    @staticmethod
    def _phrase_decompress_bytes(data, phrases, expected_len):
        result = bytearray()
        i = 0
        data_bytes = bytearray(data)
        while i < len(data_bytes):
            ch = data_bytes[i]
            if ch == 0 or ch >= 16:
                result.append(ch)
                i += 1
            else:
                if i + 1 >= len(data_bytes): break
                val = ch * 256 - 256 + data_bytes[i + 1]; i += 2
                phrase_num = val // 2
                if 0 <= phrase_num < len(phrases):
                    result.extend(phrases[phrase_num].encode('latin-1', errors='replace'))
                    if val % 2 == 1: result.append(0x20)
        return bytes(result[:expected_len])

    # ── 共享后处理 ──

    @staticmethod
    def _clean_text(raw):
        """Strip non-printable chars, collapse whitespace."""
        clean = []
        for c in raw:
            if 32 <= ord(c) < 127: clean.append(c)
            elif c in '\n\r\t': clean.append(' ')
        return re.sub(r'\s+', ' ', ''.join(clean)).strip()

    # ═══════════════════════════════════════════
    #  LZ77 解压 (ring buffer, zero-init)
    # ═══════════════════════════════════════════

    @staticmethod
    def _lz77_decompress(data, max_output=16384):
        output = bytearray()
        ring = bytearray(b'\x00' * 4096)
        ring_pos = 0
        src = bytearray(data)
        src_idx = bit_buf = bits_left = 0

        while len(output) < max_output and src_idx < len(src):
            if bits_left == 0:
                if src_idx >= len(src): break
                bit_buf, src_idx, bits_left = src[src_idx], src_idx + 1, 8
            if bit_buf & 1:
                if src_idx + 1 >= len(src): break
                low, high = src[src_idx], src[src_idx + 1]; src_idx += 2
                pos_rel = low | ((high & 0xF) << 8)
                length = (high >> 4) + 3
                for _ in range(length):
                    if len(output) >= max_output: break
                    b = ring[(ring_pos - 1 - pos_rel) & 0xFFF]
                    output.append(b)
                    ring[ring_pos & 0xFFF] = b
                    ring_pos = (ring_pos + 1) & 0xFFF
                # DON'T reset bits_left — continue processing remaining bits
                # from current bit_buf (same as helpdeco: mask shifts until overflow)
            else:
                if src_idx >= len(src): break
                b = src[src_idx]; src_idx += 1
                output.append(b)
                ring[ring_pos & 0xFFF] = b
                ring_pos = (ring_pos + 1) & 0xFFF
            bit_buf >>= 1; bits_left -= 1
        return bytes(output)

    # ═══════════════════════════════════════════
    #  |TTLBTREE → 标题
    # ═══════════════════════════════════════════

    def _read_ttlbtree(self, f, internal_files):
        ttl_offset = internal_files.get('|TTLBTREE')
        if ttl_offset is None: return []
        f.seek(ttl_offset); f.read(4); f.read(4); f.read(1)
        ttl_start = ttl_offset + 9
        f.seek(ttl_start); f.read(2); f.read(2)
        ttl_psize = struct.unpack('<H', f.read(2))[0]
        f.read(16); f.read(2); f.read(2)
        ttl_root = struct.unpack('<h', f.read(2))[0]
        f.read(2); f.read(2)
        ttl_nlevels = struct.unpack('<h', f.read(2))[0]; f.read(4)

        def read_ttl_leaf(fh, n):
            r = []
            for _ in range(n):
                off = struct.unpack('<L', fh.read(4))[0]
                nb = bytearray()
                while True:
                    b = fh.read(1)
                    if b in (b'\x00', b''): break
                    nb.extend(b)
                r.append((off, nb.decode('latin-1', errors='replace')))
            return r

        ttl_pstart = ttl_start + 38
        return self._traverse_btree_leaves(f, ttl_pstart, ttl_root, ttl_nlevels, ttl_psize, read_ttl_leaf)


class WebDocumentProcessor:
    """网页文档处理器（在线文档）"""
    
    def __init__(self):
        if not BeautifulSoup:
            raise ImportError("需要 beautifulsoup4: pip install beautifulsoup4")
        if not html2text:
            raise ImportError("需要 html2text: pip install html2text")
        
        self.converter = html2text.HTML2Text()
        self.converter.ignore_links = False
        self.converter.ignore_images = True
        self.converter.body_width = 0
    
    def process_url(self, url: str, timeout: int = 30) -> Optional[Dict]:
        try:
            import requests
            
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
            }
            response = requests.get(url, headers=headers, timeout=timeout)
            response.encoding = response.apparent_encoding or 'utf-8'
            
            html_content = response.text
            soup = BeautifulSoup(html_content, 'lxml' if _has_lxml() else 'html.parser')
            
            for tag in soup(['script', 'style', 'nav', 'footer', 'header']):
                tag.decompose()
            
            title = soup.title.get_text(strip=True) if soup.title else url
            markdown = self.converter.handle(html_content)
            text_content = soup.get_text(separator='\n', strip=True)
            
            return {
                'title': title,
                'content': markdown,
                'content_type': 'web',
                'url': url,
                'size': len(text_content),
                'line_count': text_content.count('\n') + 1,
                'hash': hashlib.md5(text_content.encode()).hexdigest(),
                'sections': self._extract_sections(soup),
                'code_examples': self._extract_code_blocks(soup)
            }
        except Exception:
            return None
    
    def _extract_sections(self, soup) -> List[Dict]:
        sections = []
        for tag in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6']):
            level = int(tag.name[1])
            title = tag.get_text(strip=True)
            sections.append({'level': level, 'title': title})
        return sections
    
    def _extract_code_blocks(self, soup) -> List[str]:
        return [tag.get_text(strip=True) for tag in soup.find_all(['pre', 'code'])]


def _has_lxml() -> bool:
    try:
        import lxml
        return True
    except ImportError:
        return False


def _process_document_worker(args: Tuple) -> List[Dict]:
    """
    处理单个文档的工作函数（用于多进程）
    返回文档信息字典列表（HLP 等格式可能返回多个文档）
    """
    file_path_str, directory_str = args
    file_path = Path(file_path_str)
    directory = Path(directory_str)
    
    processors = [
        TextProcessor(),
        MarkdownProcessor(),
        HTMLProcessor(),
        DocxProcessor(),
        DocProcessor(),
        PDFProcessor(),
        EPUBProcessor(),
        ChmProcessor(),
        HlpProcessor()
    ]
    
    try:
        stat_info = file_path.stat()
        
        for processor in processors:
            if processor.can_process(file_path):
                result = processor.process(file_path)
                if result:
                    rel_path = str(file_path.relative_to(directory)).replace('\\', '/')
                    common_meta = {
                        'full_path': str(file_path),
                        'extension': file_path.suffix.lower(),
                        'file_size': stat_info.st_size,
                        'last_modified': datetime.fromtimestamp(stat_info.st_mtime).isoformat()
                    }
                    
                    if isinstance(result, list):
                        docs = []
                        for item in result:
                            item.setdefault('path', rel_path)
                            item.update(common_meta)
                            if not item.get('path') or item['path'] == rel_path:
                                item['path'] = f"{rel_path}#{item.get('title', 'untitled')}"
                            docs.append(item)
                        return docs
                    else:
                        result['path'] = rel_path
                        result.update(common_meta)
                        return [result]
        
        return []
        
    except Exception:
        return []


class GenericDocumentScanner:
    """通用文档扫描器"""
    
    SUPPORTED_EXTENSIONS = [
        '.txt', '.md', '.markdown',
        '.htm', '.html',
        '.docx', '.doc',
        '.pdf', '.epub',
        '.hlp', '.chm'
    ]
    
    def __init__(self, kb_dir: str, config: Optional[Dict] = None, progress_callback: Optional[Callable] = None):
        self.kb_dir = Path(kb_dir)
        self.config = config or self._load_config()
        self.progress_callback = progress_callback
        
        db_file = self.config.get('database', {}).get('file', 'documents.sqlite')
        self.db_path = self.kb_dir / db_file
        
        self._scanning = False
        self._scan_thread: Optional[threading.Thread] = None
        
        # 初始化 FTS5 懒加载管理器（需要在 _init_database 之前）
        self.fts5_manager = FTS5LazyManager(
            db_path=str(self.db_path),
            main_table='documents',
            fts_table='documents_fts',
            columns=['title', 'content']
        )
        
        self._init_database()
    
    @staticmethod
    def _detect_language(title: str, content: str = '') -> str:
        """
        检测文档语言
        
        Args:
            title: 文档标题
            content: 文档内容
        
        Returns:
            语言代码: 'zh' (中文), 'ja' (日文), 'ko' (韩文), 'en' (英文), 'other' (其他)
        """
        import re
        
        # 合并标题和内容前1000字符进行检测
        sample = (title + ' ' + content[:1000])
        
        # 统计各语言字符数
        zh_chars = len(re.findall(r'[\u4e00-\u9fff]', sample))
        ja_chars = len(re.findall(r'[\u3040-\u309f\u30a0-\u30ff]', sample))
        ko_chars = len(re.findall(r'[\uac00-\ud7af]', sample))
        en_chars = len(re.findall(r'[A-Za-z]', sample))
        
        # 选择字符数最多的语言
        counts = {'zh': zh_chars, 'ja': ja_chars, 'ko': ko_chars, 'en': en_chars}
        max_lang = max(counts, key=counts.get)
        
        # 如果最大字符数太少，返回 other
        if counts[max_lang] < 10:
            return 'other'
        
        return max_lang
    
    @staticmethod
    def _apply_performance_pragmas(conn, for_build=False):
        """应用 SQLite 性能 PRAGMA 设置（镜像 sqlite_vector_query_knowledge_base.py 的配置）"""
        conn.execute("PRAGMA journal_mode=WAL")
        conn.execute("PRAGMA synchronous=OFF" if for_build else "PRAGMA synchronous=NORMAL")
        conn.execute("PRAGMA cache_size=-256000")  # 256MB 缓存
        conn.execute("PRAGMA temp_store=MEMORY")
        conn.execute("PRAGMA mmap_size=268435456")  # 256MB 内存映射
        conn.execute("PRAGMA busy_timeout=10000")

    @staticmethod
    def _find_7zip_path() -> Optional[str]:
        """查找 7-Zip 可执行文件路径"""
        tools_dir = Path(__file__).parent.parent.parent.parent / 'tools'
        candidates = [
            str(tools_dir / '7z' / '7z.exe'),
            r'C:\Program Files\7-Zip\7z.exe',
            r'C:\Program Files (x86)\7-Zip\7z.exe',
        ]
        for p in candidates:
            if Path(p).exists():
                return p
        return None

    @staticmethod
    def _estimate_total_docs(all_files: List[Path]) -> int:
        """估算文档扫描的总文档数（展开 CHM 内子文档数），用于准确进度"""
        estimated = 0
        sevenzip = GenericDocumentScanner._find_7zip_path()
        for f in all_files:
            suffix = f.suffix.lower()
            if suffix == '.chm' and sevenzip:
                try:
                    result = subprocess.run(
                        [sevenzip, 'l', '-ba', str(f)],
                        capture_output=True, text=True, timeout=60,
                        creationflags=getattr(subprocess, 'CREATE_NO_WINDOW', 0)
                    )
                    if result.returncode == 0:
                        html_count = sum(1 for line in result.stdout.strip().split('\n') if line.strip())
                        estimated += max(html_count, 1)
                    else:
                        estimated += 5000
                except Exception:
                    estimated += 5000
            elif suffix in ('.chm', '.hlp', '.epub'):
                estimated += 5000
            else:
                estimated += 1
        return max(estimated, 1)

    def _load_config(self) -> Dict:
        """加载配置文件"""
        config_path = self.kb_dir / "config.json"
        if config_path.exists():
            try:
                with open(config_path, 'r', encoding='utf-8') as f:
                    return json.load(f)
            except Exception:
                pass
        
        return {
            'database': {'file': 'documents.sqlite'},
            'build': {
                'parallel_workers': None,
                'batch_size': 50,
                'supported_extensions': self.SUPPORTED_EXTENSIONS
            }
        }
    
    def save_config(self):
        """保存配置到文件"""
        config_path = self.kb_dir / "config.json"
        config_path.parent.mkdir(parents=True, exist_ok=True)
        
        default_config = {
            'name': 'document-knowledge-base',
            'type': 'generic-documents',
            'version': '1.0',
            'database': {
                'file': 'documents.sqlite'
            },
            'build': {
                'parallel_workers': None,
                'batch_size': 50,
                'supported_extensions': self.SUPPORTED_EXTENSIONS
            }
        }
        
        merged = {**default_config, **self.config}
        
        with open(config_path, 'w', encoding='utf-8') as f:
            json.dump(merged, f, indent=2, ensure_ascii=False)
    
    def _init_database(self):
        conn = sqlite3.connect(str(self.db_path))
        self._apply_performance_pragmas(conn)
        cursor = conn.cursor()
        
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS documents (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                path TEXT NOT NULL,
                full_path TEXT NOT NULL,
                extension TEXT,
                title TEXT,
                title_lower TEXT,
                title_rev TEXT,
                content TEXT,
                content_type TEXT,
                file_size INTEGER,
                size INTEGER,
                line_count INTEGER,
                hash TEXT,
                last_modified TEXT,
                sections TEXT,
                code_examples TEXT,
                url TEXT,
                requires_extraction INTEGER DEFAULT 0
            )
        """)
        
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS document_entities (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                document_id INTEGER NOT NULL,
                name TEXT NOT NULL,
                kind TEXT NOT NULL,
                line INTEGER,
                definition TEXT,
                FOREIGN KEY (document_id) REFERENCES documents(id)
            )
        """)
        
        cursor.execute("CREATE INDEX IF NOT EXISTS idx_documents_path ON documents(path)")
        cursor.execute("CREATE INDEX IF NOT EXISTS idx_documents_hash ON documents(hash)")
        cursor.execute("CREATE INDEX IF NOT EXISTS idx_document_entities_name ON document_entities(name)")
        
        # Schema 迁移：添加 title_lower 和 title_rev 字段
        cursor.execute("PRAGMA table_info(documents)")
        columns = [row[1] for row in cursor.fetchall()]
        
        if 'title_lower' not in columns:
            cursor.execute("ALTER TABLE documents ADD COLUMN title_lower TEXT")
        
        if 'title_rev' not in columns:
            cursor.execute("ALTER TABLE documents ADD COLUMN title_rev TEXT")
        
        if 'language' not in columns:
            cursor.execute("ALTER TABLE documents ADD COLUMN language TEXT DEFAULT 'en'")
        
        # 创建逆序索引
        cursor.execute("CREATE INDEX IF NOT EXISTS idx_documents_title_lower ON documents(title_lower)")
        cursor.execute("CREATE INDEX IF NOT EXISTS idx_documents_title_rev ON documents(title_rev)")
        cursor.execute("CREATE INDEX IF NOT EXISTS idx_documents_language ON documents(language)")
        
        # 填充已有数据的逆序字段和语言字段
        cursor.execute("SELECT COUNT(*) FROM documents WHERE title IS NOT NULL AND title_rev IS NULL")
        missing = cursor.fetchone()[0]
        if missing > 0:
            conn.create_function("my_reverse", 1, lambda s: s[::-1].lower() if s else '')
            conn.create_function("my_lower", 1, lambda s: s.lower() if s else '')
            cursor.execute("UPDATE documents SET title_lower = my_lower(title), title_rev = my_reverse(title) WHERE title IS NOT NULL AND title_rev IS NULL")
        
        # 填充已有数据的语言字段（仅处理 language IS NULL 的记录）
        cursor.execute("SELECT COUNT(*) FROM documents WHERE language IS NULL")
        missing_lang = cursor.fetchone()[0]
        if missing_lang > 0 and missing_lang < 50000:
            cursor.execute("SELECT id, title, content FROM documents WHERE language IS NULL")
            for row in cursor.fetchall():
                doc_id, title, content = row
                lang = self._detect_language(title or '', content or '')
                cursor.execute("UPDATE documents SET language = ? WHERE id = ?", (lang, doc_id))
        elif missing_lang >= 50000:
            logger.warning(f"跳过 {missing_lang} 条语言检测（数量过大），新插入文档会自动检测语言")
        
        # 创建 FTS5 虚拟表（懒加载，不填充数据）
        self.fts5_manager.create_fts_table(conn)
        
        conn.commit()
        conn.close()
    
    def scan_directory(self, directory: str, extensions: Optional[List[str]] = None,
                      max_workers: Optional[int] = None,
                      exclude_dirs: Optional[List[str]] = None) -> Dict:
        """
        扫描目录中的文档
        
        Args:
            directory: 要扫描的目录
            extensions: 要处理的文件扩展名列表
            max_workers: 最大工作进程数
            exclude_dirs: 要排除的子目录名列表（默认排除多语言帮助子目录）
        
        Returns:
            扫描统计信息
        """
        directory = Path(directory)
        if not directory.exists():
            return {'error': f'目录不存在: {directory}'}
        
        extensions = extensions or self.config.get('build', {}).get('supported_extensions', self.SUPPORTED_EXTENSIONS)
        extensions = [e.lower() for e in extensions]
        
        exclude_set = set(exclude_dirs) if exclude_dirs else DEFAULT_EXCLUDE_DIRS
        
        all_files = []
        for ext in extensions:
            for f in directory.rglob(f'*{ext}'):
                # 排除语言子目录（如 ja/、fr/、de/ 等）
                if exclude_set.intersection(f.relative_to(directory).parts):
                    continue
                all_files.append(f)
        
        if not all_files:
            return {'total_files': 0, 'processed': 0, 'failed': 0}
        
        total_files = len(all_files)
        estimated_total = self._estimate_total_docs(all_files)
        
        parallel_workers_config = self.config.get('build', {}).get('parallel_workers')
        if max_workers is None:
            if parallel_workers_config:
                max_workers = max(1, parallel_workers_config)
            else:
                cpu_cores = cpu_count()
                # 至少 2 个 worker（帮助知识库策略: max(2, cpu_cores-1)），不超过文件数
                max_workers = min(max(2, cpu_cores - 1), total_files)
        
        batch_size_config = self.config.get('build', {}).get('batch_size', 50)
        # chunksize 不能超过平均每 worker 文件数，否则重文件会堆积
        files_per_worker = max(1, total_files // max_workers) if max_workers > 0 else 50
        chunk_size = min(max(1, files_per_worker), max(20, batch_size_config))
        
        processed = 0
        failed = 0
        warnings = []
        install_hints = []
        
        # 检测 PDF 文件依赖
        if '.pdf' in extensions:
            has_pymupdf = False
            has_pdfplumber = False
            try:
                import fitz
                has_pymupdf = True
            except ImportError:
                pass
            try:
                import pdfplumber
                has_pdfplumber = True
            except ImportError:
                pass
            
            pdf_files = [f for f in all_files if f.suffix.lower() == '.pdf']
            if pdf_files and not (has_pymupdf or has_pdfplumber):
                warnings.append(f"发现 {len(pdf_files)} 个 PDF 文件，但缺少依赖（PyMuPDF 或 pdfplumber）")
                install_hints.append({
                    'type': 'pdf',
                    'message': 'PDF 处理需要安装依赖（推荐优先使用 PyMuPDF）',
                    'commands': ['pip install PyMuPDF', 'pip install pdfplumber'],
                    'affected_files': len(pdf_files)
                })
        
        # 检测 DOCX 文件依赖
        if '.docx' in extensions:
            has_docx = False
            try:
                from docx import Document
                has_docx = True
            except ImportError:
                pass
            
            docx_files = [f for f in all_files if f.suffix.lower() == '.docx']
            if docx_files and not has_docx:
                warnings.append(f"发现 {len(docx_files)} 个 DOCX 文件，但缺少依赖（python-docx）")
                install_hints.append({
                    'type': 'docx',
                    'message': 'Word .docx 处理需要安装依赖',
                    'commands': ['pip install python-docx'],
                    'affected_files': len(docx_files)
                })
        
        # 检测 DOC 文件依赖（需要系统工具）
        if '.doc' in extensions:
            doc_files = [f for f in all_files if f.suffix.lower() == '.doc']
            if doc_files:
                has_antiword = shutil.which('antiword') is not None
                has_catdoc = shutil.which('catdoc') is not None
                
                if not (has_antiword or has_catdoc):
                    warnings.append(f"发现 {len(doc_files)} 个 DOC 文件，但缺少工具（antiword 或 catdoc）")
                    install_hints.append({
                        'type': 'doc',
                        'message': 'Word .doc 处理需要安装系统工具',
                        'commands': ['winget install antiword', 'choco install catdoc', '或从官网下载'],
                        'affected_files': len(doc_files),
                        'note': '这是系统级工具，不是 Python 包'
                    })
        
        conn = sqlite3.connect(str(self.db_path))
        self._apply_performance_pragmas(conn, for_build=True)
        cursor = conn.cursor()
        
        existing_mtimes = {}
        try:
            cursor.execute("SELECT full_path, last_modified FROM documents")
            for r in cursor.fetchall():
                existing_mtimes[r[0]] = r[1]
        except Exception:
            pass
        
        changed = []
        for fx in all_files:
            try:
                stx = fx.stat()
                mt = datetime.fromtimestamp(stx.st_mtime).isoformat()
                if existing_mtimes.get(str(fx)) == mt:
                    continue
                changed.append(fx)
            except Exception:
                changed.append(fx)
        
        skipped = total_files - len(changed)
        all_files = changed
        
        if not all_files:
            conn.close()
            return {'total_files': total_files, 'processed': 0, 'failed': 0,
                    'skipped': skipped, 'warnings': warnings, 'install_hints': install_hints}
        
        for fx in all_files:
            try:
                # 先获取要删除的文档ID
                cursor.execute("SELECT id FROM documents WHERE full_path = ?", (str(fx),))
                doc_ids = [row[0] for row in cursor.fetchall()]
                
                # 删除documents表记录
                cursor.execute("DELETE FROM documents WHERE full_path = ?", (str(fx),))
                
                # 同步删除FTS5索引
                for doc_id in doc_ids:
                    cursor.execute("DELETE FROM documents_fts WHERE rowid = ?", (doc_id,))
            except Exception:
                pass
        conn.commit()
        
        try:
            COMMIT_INTERVAL = 500  # 每处理多少文档提交一次
            batch_count = 0
            
            with ProcessPoolExecutor(max_workers=max_workers) as executor:
                args = [(str(f), str(directory)) for f in all_files]
                future_map = {executor.submit(_process_document_worker, a): a for a in args}
                
                file_idx = 0
                for future in as_completed(future_map):
                    try:
                        docs = future.result()
                    except Exception:
                        failed += 1
                        file_idx += 1
                        continue
                    
                    if not docs:
                        file_idx += 1
                        continue
                    
                    for result in docs:
                        if result.get('requires_dependencies'):
                            failed += 1
                            continue
                        
                        try:
                            title = result.get('title', '')
                            title_lower = title.lower() if title else ''
                            title_rev = title[::-1].lower() if title else ''
                            content = result.get('content', '')
                            language = self._detect_language(title, content)
                            
                            # 清理 Base64 data URI（html2text ignore_images=True 已处理，此处做防御性清理）
                            content = re.sub(
                                r'data:image/[^;]+;base64,[A-Za-z0-9+/=]{100,}',
                                '[Base64 Image]',
                                content
                            )
                            requires_extraction = result.get('requires_extraction', 0)
                            
                            cursor.execute("""
                                INSERT INTO documents (
                                    path, full_path, extension, title, title_lower, title_rev,
                                    content, content_type, file_size, size, line_count,
                                    hash, last_modified, sections, code_examples, url, requires_extraction, language
                                ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                            """, (
                                result.get('path'),
                                result.get('full_path'),
                                result.get('extension'),
                                title,
                                title_lower,
                                title_rev,
                                content,
                                result.get('content_type'),
                                result.get('file_size'),
                                result.get('size'),
                                result.get('line_count'),
                                result.get('hash'),
                                result.get('last_modified'),
                                json.dumps(result.get('sections', [])),
                                json.dumps(result.get('code_examples', [])),
                                result.get('url'),
                                requires_extraction,
                                language
                            ))
                            processed += 1
                            batch_count += 1
                        except Exception:
                            failed += 1
                    
                    # 每完成一个文件或达到提交间隔，就 commit 一次
                    if batch_count >= COMMIT_INTERVAL:
                        conn.commit()
                        batch_count = 0
                    
                    file_idx += 1
                    if self.progress_callback:
                        self.progress_callback(processed, estimated_total)
            
            conn.commit()
            
            # 数据库维护：回收碎片、优化查询计划
            if processed > 0:
                if self.progress_callback:
                    self.progress_callback(total_files, "正在优化数据库（VACUUM）...")
                conn.execute("PRAGMA incremental_vacuum")
                conn.execute("VACUUM")
                conn.execute("PRAGMA optimize")
                conn.execute("ANALYZE")
                
                # 构建完成后全量构建 FTS5 索引
                if self.progress_callback:
                    self.progress_callback(total_files, "正在构建 FTS5 全文索引...")
                try:
                    self.fts5_manager.rebuild_full()
                except Exception as e:
                    logger.warning(f"FTS5 索引构建失败（后续查询自动降级）: {e}")
            
        finally:
            # 恢复 synchronous 为非构建模式
            try:
                conn.execute("PRAGMA synchronous=NORMAL")
            except Exception:
                pass
            conn.close()
        
        return {
            'total_files': total_files,
            'processed': processed,
            'failed': failed,
            'skipped': skipped,
            'warnings': warnings,
            'install_hints': install_hints
        }
    
    def scan_directory_async(self, directory: str, extensions: Optional[List[str]] = None,
                            max_workers: Optional[int] = None, callback: Optional[Callable] = None,
                            exclude_dirs: Optional[List[str]] = None):
        """
        异步扫描目录中的文档（立即返回，后台处理）
        
        Args:
            directory: 要扫描的目录
            extensions: 要处理的文件扩展名列表
            max_workers: 最大工作进程数
            callback: 完成后的回调函数
        
        Returns:
            立即返回，实际处理在后台线程中进行
        """
        if self._scanning:
            print("扫描已在进行中...")
            return
        
        def _scan_task():
            try:
                self._scanning = True
                result = self.scan_directory(directory, extensions, max_workers, exclude_dirs)
                if callback:
                    callback(result)
                if self.progress_callback:
                    self.progress_callback(100, "扫描完成")
            finally:
                self._scanning = False
        
        self._scan_thread = threading.Thread(target=_scan_task, daemon=True)
        self._scan_thread.start()
        
        print(f"异步扫描已启动: {directory}")
    
    def is_scanning(self) -> bool:
        """检查是否正在扫描"""
        return self._scanning
    
    def wait_scan_complete(self, timeout: Optional[float] = None) -> bool:
        """
        等待扫描完成
        
        Args:
            timeout: 超时时间（秒）
        
        Returns:
            是否完成（True=完成，False=超时）
        """
        if self._scan_thread is None:
            return True
        
        self._scan_thread.join(timeout=timeout)
        return not self._scan_thread.is_alive()
    
    def add_web_document(self, url: str) -> Optional[Dict]:
        """
        添加网页文档
        
        Args:
            url: 网页 URL
        
        Returns:
            文档信息或 None
        """
        try:
            processor = WebDocumentProcessor()
            result = processor.process_url(url)
            
            if result:
                conn = sqlite3.connect(str(self.db_path))
                self._apply_performance_pragmas(conn)
                cursor = conn.cursor()
                
                try:
                    title = result.get('title', '')
                    content = result.get('content', '')
                    # 防御性清理 Base64 data URI
                    content = re.sub(
                        r'data:image/[^;]+;base64,[A-Za-z0-9+/=]{100,}',
                        '[Base64 Image]',
                        content
                    )
                    title_lower = title.lower() if title else ''
                    title_rev = title[::-1].lower() if title else ''
                    language = self._detect_language(title, content)
                    
                    cursor.execute("""
                        INSERT INTO documents (
                            path, full_path, extension, title, title_lower, title_rev, content, content_type,
                            file_size, size, line_count, hash, last_modified,
                            sections, code_examples, url, language
                        ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                    """, (
                        url,
                        url,
                        '.html',
                        title,
                        title_lower,
                        title_rev,
                        content,
                        result.get('content_type'),
                        result.get('size'),
                        result.get('size'),
                        result.get('line_count'),
                        result.get('hash'),
                        datetime.now().isoformat(),
                        json.dumps(result.get('sections', [])),
                        json.dumps(result.get('code_examples', [])),
                        result.get('url'),
                        language
                    ))
                    
                    conn.commit()
                finally:
                    conn.close()
                
                return result
            
            return None
        except Exception as e:
            return {'error': str(e)}
    
    def search(self, query: str, content_type: Optional[str] = None, top_k: int = 10, use_fts5: bool = True) -> List[Dict]:
        """
        搜索文档（支持 FTS5 懒加载 + 逆序索引优化）
        
        Args:
            query: 搜索关键词（支持空格分隔的多关键词）
            content_type: 文档类型过滤
            top_k: 返回结果数
            use_fts5: 是否使用 FTS5（默认 True）
        
        Returns:
            匹配的文档列表（按相关性评分排序）
        """
        if use_fts5:
            # 使用 FTS5 懒加载搜索（自动降级 + 后台构建）
            results = self.fts5_manager.search(
                query=query,
                search_func=lambda q: self._reverse_index_search(q, content_type, top_k),
                top_k=top_k,
                use_bM25=True
            )
        else:
            # 直接使用逆序索引搜索
            results = self._reverse_index_search(query, content_type, top_k)
        
        # 后置 content_type 过滤（弥补 FTS5 _fts_search 不支持 content_type 过滤的问题）
        if content_type and results:
            results = [r for r in results if r.get('content_type') == content_type]
        return results
    
    def _reverse_index_search(self, query: str, content_type: Optional[str] = None, top_k: int = 10) -> List[Dict]:
        """
        逆序索引搜索（降级搜索）
        
        Args:
            query: 搜索关键词
            content_type: 文档类型过滤
            top_k: 返回结果数
        
        Returns:
            匹配的文档列表
        """
        conn = sqlite3.connect(str(self.db_path))
        self._apply_performance_pragmas(conn)
        conn.row_factory = sqlite3.Row
        cursor = conn.cursor()
        
        try:
            keywords = query.split()
            if not keywords:
                return []
            
            # 使用逆序索引优化标题搜索
            title_conditions = []
            content_conditions = []
            score_parts = []
            params = []
            
            for kw in keywords:
                kw_lower = kw.lower()
                kw_rev = kw_lower[::-1]
                
                # 标题搜索：优先使用逆序索引
                title_conditions.append("title_lower = ?")
                title_conditions.append("title_lower GLOB ?")
                title_conditions.append("title_rev GLOB ?")
                title_conditions.append("title_lower LIKE ?")
                params.extend([kw_lower, f'{kw_lower}*', f'{kw_rev}*', f'%{kw_lower}%'])
                
                # 内容搜索
                content_conditions.append("content LIKE ?")
                params.append(f'%{kw}%')
                
                # 评分
                score_parts.append(f"(CASE WHEN title_lower = ? THEN 20 ELSE 0 END)")
                score_parts.append(f"(CASE WHEN title_lower GLOB ? THEN 15 ELSE 0 END)")
                score_parts.append(f"(CASE WHEN title_rev GLOB ? THEN 15 ELSE 0 END)")
                score_parts.append(f"(CASE WHEN title_lower LIKE ? THEN 10 ELSE 0 END)")
                score_parts.append(f"(CASE WHEN content LIKE ? THEN 1 ELSE 0 END)")
                params.extend([kw_lower, f'{kw_lower}*', f'{kw_rev}*', f'%{kw_lower}%', f'%{kw}%'])
            
            title_where = " OR ".join(title_conditions)
            content_where = " OR ".join(content_conditions)
            score_expr = " + ".join(score_parts)
            
            if content_type:
                sql = f"""
                    SELECT *, ({score_expr}) as score
                    FROM documents
                    WHERE content_type = ? AND ({title_where} OR {content_where})
                    ORDER BY score DESC, size DESC
                    LIMIT ?
                """
                params.insert(0, content_type)
            else:
                sql = f"""
                    SELECT *, ({score_expr}) as score
                    FROM documents
                    WHERE {title_where} OR {content_where}
                    ORDER BY score DESC, size DESC
                    LIMIT ?
                """
            
            params.append(top_k)
            cursor.execute(sql, params)
            
            results = []
            for row in cursor.fetchall():
                results.append(dict(row))
            
            return results
        finally:
            conn.close()
    
    def crawl_website(self, start_url: str, max_pages: int = 100, max_depth: int = 3,
                     domain_filter: Optional[str] = None, url_pattern: Optional[str] = None) -> Dict:
        """
        自动爬取网站（递归发现链接）
        
        Args:
            start_url: 起始 URL
            max_pages: 最大页面数
            max_depth: 最大深度
            domain_filter: 域名过滤（只爬取该域名下的页面）
            url_pattern: URL 正则模式过滤
        
        Returns:
            爬取统计信息
        """
        try:
            import requests
            from urllib.parse import urljoin, urlparse
            from collections import deque
        except ImportError:
            return {'error': '需要安装 requests: pip install requests'}
        
        visited = set()
        queue = deque([(start_url, 0)])
        stats = {'success': 0, 'failed': 0, 'skipped': 0, 'total_size': 0}
        
        processor = WebDocumentProcessor()
        
        while queue and stats['success'] < max_pages:
            url, depth = queue.popleft()
            
            if url in visited:
                continue
            
            if depth > max_depth:
                continue
            
            visited.add(url)
            
            # 进度回调
            if self.progress_callback and stats['success'] % 10 == 0:
                self.progress_callback(stats['success'], f"已爬取 {stats['success']} 个页面")
            
            # 处理当前页面
            result = processor.process_url(url, timeout=30)
            
            if result:
                # 添加到知识库
                self.add_web_document(url)
                stats['success'] += 1
                stats['total_size'] += result.get('size', 0)
                
                # 提取新链接
                if depth < max_depth:
                    try:
                        resp = requests.get(url, timeout=20, headers={
                            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
                        })
                        html_content = resp.text
                        
                        new_links = self._extract_links_from_html(
                            html_content, url,
                            domain_filter=domain_filter,
                            url_pattern=url_pattern
                        )
                        
                        for link in new_links:
                            if link not in visited:
                                queue.append((link, depth + 1))
                    except Exception:
                        pass
            else:
                stats['failed'] += 1
        
        if self.progress_callback:
            self.progress_callback(100, f"爬取完成: {stats['success']} 个页面")
        
        return stats
    
    def _extract_links_from_html(self, html_content: str, base_url: str,
                                 domain_filter: Optional[str] = None,
                                 url_pattern: Optional[str] = None) -> List[str]:
        """从 HTML 中提取链接"""
        if not BeautifulSoup:
            return []
        
        from urllib.parse import urljoin, urlparse
        
        soup = BeautifulSoup(html_content, 'html.parser')
        links = []
        
        for tag in soup.find_all('a', href=True):
            href = tag['href']
            full_url = urljoin(base_url, href)
            parsed = urlparse(full_url)
            
            # 移除锚点
            if '#' in full_url:
                full_url = full_url.split('#')[0]
            
            # 域名过滤
            if domain_filter and parsed.netloc != domain_filter:
                continue
            
            # URL 模式过滤
            if url_pattern:
                if not re.search(url_pattern, full_url):
                    continue
            
            # 只保留 HTML 页面
            if not full_url.endswith(('.html', '.htm', '/')):
                continue
            
            # 忽略非 HTML 文件
            if any(ext in full_url.lower() for ext in ['.pdf', '.zip', '.png', '.jpg', '.gif', '.css', '.js']):
                continue
            
            links.append(full_url)
        
        return list(set(links))
    
    def get_statistics(self) -> Dict:
        """获取统计信息"""
        conn = sqlite3.connect(str(self.db_path))
        self._apply_performance_pragmas(conn)
        cursor = conn.cursor()
        
        try:
            cursor.execute("SELECT COUNT(*) FROM documents")
            total_documents = cursor.fetchone()[0]
            
            cursor.execute("""
                SELECT content_type, COUNT(*) 
                FROM documents 
                GROUP BY content_type
            """)
            by_type = dict(cursor.fetchall())
            
            # 直接从数据库查询扩展名分布（替代不存在的 language 列）
            cursor.execute("""
                SELECT COALESCE(extension, '(no ext)') AS ext, COUNT(*) AS cnt
                FROM documents
                GROUP BY ext
                ORDER BY cnt DESC
            """)
            by_extension = dict(cursor.fetchall())
            
            # 实际数据库文件磁盘大小
            db_size_mb = self.db_path.stat().st_size / (1024 * 1024)
            
            return {
                'total_documents': total_documents,
                'by_type': by_type,
                'by_extension': by_extension,
                'database_size_mb': round(db_size_mb, 2),
            }
        finally:
            conn.close()
